from docx import Document
from docx.shared import Inches
import pyttsx3

documento = Document()

def speak(text):
    pyttsx3.speak(text)

# profile picture
documento.add_picture(
    'project.png',
    width=Inches(2.0)
)

# informacion personal
nombre = input('¿Cuál es tu nombre y apellido?: ')
telefono = input('¿Cuál es tu número de teléfono?')
email = input('¿Cuál es tu email?')

documento.add_heading('Información personal')

informacion_personal = documento.add_paragraph()
informacion_personal.add_run('Nombre y apellido: ').bold = True
informacion_personal.add_run(nombre + '\n')
informacion_personal.add_run('Teléfono: ').bold = True
informacion_personal.add_run(telefono + '\n')
informacion_personal.add_run('Email: ').bold = True
informacion_personal.add_run(email + '\n')

# sobre mi
documento.add_heading('Sobre mí')

speak('Ahora, dime informacion valiosa sobre ti.')
sobre_mi = input('Cuentame sobre tí, tu perfil, tus aspiraciones ... ')

documento.add_paragraph(sobre_mi)

# formacion

speak('Estamos en el area de formación, cuentame que has estudiado y donde lo hiciste.')
documento.add_heading('Formación')

p_formacion = documento.add_paragraph()

tiene_formacion = True

while tiene_formacion:
    respuesta = input('¿Tienes formacion para agregar? (Si/No): ')
    if (respuesta.upper() == 'SI'):
        lugar = input('Lugar de estudio: ')
        desde_fecha = input('Dime el año de ingreso: ')
        hasta_fecha = input('Dime el año de egreso: ')
        titulo = input('¿Cuál es el titulo de la formación?: ')
        p_formacion.add_run(lugar + ': ').bold = True
        p_formacion.add_run(desde_fecha + '-' + hasta_fecha + '\n').italic = True
        p_formacion.add_run(titulo + '\n')
    else:
        tiene_formacion = False

# intereses
speak('Cuentame a cerca de tus intereses.')

documento.add_heading('Intereses')

p_intereses = documento.add_paragraph()
p_intereses.style = 'List Bullet'

tiene_intereses = True

while tiene_intereses:
    respuesta_i = input('¿Tienes intereses para agregar? (Si/No): ')
    if (respuesta_i.upper() == 'SI'):
        interes = input('Dime un interes que tengas: ')
        p_intereses.add_run(interes + '\n')
    else:
        tiene_intereses = False

# finalizacion
documento.save('curriculumgenerado.docx')

speak('Tu curriculum ha sido generado, busca en tu repositorio.')






